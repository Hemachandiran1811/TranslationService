from fastapi import FastAPI, UploadFile, File, HTTPException,Form
from fastapi.middleware.cors import CORSMiddleware
import google.generativeai as genai
from PyPDF2 import PdfReader
from docx import Document
from fastapi.responses import StreamingResponse
from fpdf import FPDF

from flask_cors import CORS
from flask import Flask
app = FastAPI()
app = Flask(__name__)
CORS(app)
app = FastAPI()
# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allows all origins, you can specify a list of allowed origins instead of "*"
    allow_credentials=True,
    allow_methods=["*"],  # Allows all methods (POST, GET, etc.)
    allow_headers=["*"],  # Allows all headers
)

# Configure the Google AI SDK
genai.configure(api_key="AIzaSyCXMNUqhnPX_EV1BPWgfs0oqGcAybwozzA")

generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 64,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

model = genai.GenerativeModel(
    model_name="gemini-1.5-flash",
    generation_config=generation_config
)

@app.post("/upload-summarize-translate/")
async def upload_summarize_translate(language: str = Form(...), 
                                     file: UploadFile = File(...), 
                                     output_format: str = Form(...)):
    try:
        print("language received: ", language)
        print("File received: ", file)
        print("Output format: ", output_format)

        # Determine file type and extract text
        file_extension = file.filename.split(".")[-1].lower()
        text = ""

        if file_extension == "pdf":
            reader = PdfReader(file.file)
            for page in reader.pages:
                text += page.extract_text()
        elif file_extension == "txt":
            text = (await file.read()).decode("utf-8")
        elif file_extension == "docx":
            doc = Document(file.file)
            for para in doc.paragraphs:
                text += para.text
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format")

        # Limit the extracted text to 400 words
        words = text.split()
        limited_text = " ".join(words[:400])

        # Prepare the prompt for summarization
        prompt = f"summarize the following words in {language}: {limited_text}"

        # Start a chat session with the Gemini LLM
        chat_session = model.start_chat(history=[
            {
                "role": "user",
                "parts": [prompt],
            }
        ])

        # Get the summary from the model
        summary_response = chat_session.send_message(prompt)
        summary = summary_response.text

        # Remove newline characters from the summary
        summary = summary.replace('\n', ' ')

        # Save the summary in the desired file format
        output_file_path = ""
        if output_format.lower() == "txt":
            output_file_path = "summary.txt"
            with open(output_file_path, "w", encoding="utf-8") as f:
                f.write(summary)
        elif output_format.lower() == "docx":
            output_file_path = "summary.docx"
            doc = Document()
            doc.add_paragraph(summary)
            doc.save(output_file_path)
        elif output_format.lower() == "pdf":
            output_file_path = "summary.pdf"
            
            # Use fpdf2 and set the path to the Unicode font
            pdf = FPDF()
            pdf.add_page()
            pdf.add_font("DejaVu", '', './fonts/DejaVuSansCondensed.ttf', uni=True)  # Provide the correct path
            pdf.set_font("DejaVu", size=12)
            pdf.multi_cell(0, 10, summary)
            pdf.output(output_file_path)
        else:
            raise HTTPException(status_code=400, detail="Unsupported output format")

        # Return the saved file as a binary stream
        def iterfile():
            with open(output_file_path, "rb") as f:
                yield from f

        return StreamingResponse(iterfile(), 
                                 media_type="application/octet-stream", 
                                 headers={
                                     "Content-Disposition": f"attachment; filename={output_file_path}"
                                 })

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
